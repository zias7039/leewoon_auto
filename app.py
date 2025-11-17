import io
import os
import re
import tempfile
import subprocess
from datetime import datetime, date
from decimal import Decimal
from zipfile import ZipFile, ZIP_DEFLATED, BadZipFile
from typing import Optional, Set

import streamlit as st
from openpyxl import load_workbook, Workbook
from openpyxl.utils.exceptions import InvalidFileException
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

# 스타일
from ui_style import inject as inject_style, h4, small_note

# 선택: docx2pdf
try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

# -------- 치환 유틸 --------
TOKEN_RE = re.compile(r"\{\{([A-Z]+[0-9]+)(?:\|([^}]+))?\}\}")
LEFTOVER_RE = re.compile(r"\{\{[^}]+\}\}")
DEFAULT_OUT = f"{datetime.today():%Y%m%d}_#_납입요청서_DB저축은행.docx"
TARGET_SHEET = "2.  배정후 청약시"


def ensure_docx(name: str) -> str:
    name = (name or "").strip()
    return name if name.lower().endswith(".docx") else (name + ".docx")


def ensure_pdf(name: str) -> str:
    """출력 파일명을 PDF 확장자로 정리."""
    base = (name or "output").strip()
    return base if base.lower().endswith(".pdf") else (base + ".pdf")


def has_soffice() -> bool:
    """LibreOffice(soffice) 사용 가능 여부 확인."""
    try:
        subprocess.run(
            ["soffice", "--version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=False,
        )
        return True
    except FileNotFoundError:
        return False


def try_format_as_date(v) -> str:
    try:
        if isinstance(v, (datetime, date)):
            return f"{v.year}. {v.month}. {v.day}."
        if isinstance(v, str):
            s = v.strip()
            # 2024-01-01 같은 형식만 간단히 처리
            if re.fullmatch(r"\d{4}-\d{2}-\d{2}", s):
                dt = datetime.strptime(s, "%Y-%m-%d").date()
                return f"{dt.year}. {dt.month}. {dt.day}."
    except Exception:
        pass
    return ""


def fmt_number(v) -> str:
    try:
        if isinstance(v, (int, float, Decimal)):
            return f"{float(v):,.0f}"
        if isinstance(v, str):
            raw = v.replace(",", "")
            if re.fullmatch(r"-?\d+(\.\d+)?", raw):
                return f"{float(raw):,.0f}"
    except Exception:
        pass
    return ""


def value_to_text(v) -> str:
    s = try_format_as_date(v)
    if s:
        return s
    s = fmt_number(v)
    if s:
        return s
    return "" if v is None else str(v)


def apply_inline_format(value, fmt: Optional[str]) -> str:
    if fmt is None or fmt.strip() == "":
        return value_to_text(value)

    # 날짜 포맷 (YYYY/MM/DD 등)
    if any(tok in fmt for tok in ("YYYY", "MM", "DD")):
        if isinstance(value, str) and re.fullmatch(r"\d{4}-\d{2}-\d{2}", value.strip()):
            value = datetime.strptime(value.strip(), "%Y-%m-%d").date()
        if isinstance(value, (datetime, date)):
            f = (
                fmt.replace("YYYY", "%Y")
                .replace("MM", "%m")
                .replace("DD", "%d")
            )
            return value.strftime(f)
        return value_to_text(value)

    # 숫자 포맷 (#,###.00 같은 형태)
    if re.fullmatch(r"[#,0]+(?:\.[0#]+)?", fmt.replace(",", "")):
        try:
            num = float(str(value).replace(",", ""))
            decimals = len(fmt.split(".")[1]) if "." in fmt else 0
            return f"{num:,.{decimals}f}"
        except Exception:
            return value_to_text(value)

    return value_to_text(value)


def replace_in_paragraph(paragraph: Paragraph, repl_func):
    """문단에 포함된 {{A1}} 토큰 치환."""
    if not paragraph.text:
        return
    new_text = repl_func(paragraph.text)
    if new_text == paragraph.text:
        return
    # 단순하게 run 구조는 무시하고 전체 텍스트 교체
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_in_table(cell: _Cell, repl_func):
    """테이블 셀 내부 문단 치환."""
    for p in cell.paragraphs:
        replace_in_paragraph(p, repl_func)
    for t in cell.tables:
        for row in t.rows:
            for c in row.cells:
                replace_in_table(c, repl_func)


def iter_block_items(parent):
    """문서/헤더/푸터/셀 안의 단락과 셀을 순회."""
    if hasattr(parent, "paragraphs") and hasattr(parent, "tables"):
        for p in parent.paragraphs:
            yield p
        for t in parent.tables:
            for row in t.rows:
                for cell in row.cells:
                    for item in iter_block_items(cell):
                        yield item


def replace_everywhere(doc: Document, repl_func):
    """본문 + 헤더/푸터 전체에 대해 토큰 치환."""
    # 본문
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            replace_in_paragraph(item, repl_func)

    # 헤더/푸터
    for section in doc.sections:
        for container in (section.header, section.footer):
            for item in iter_block_items(container):
                if isinstance(item, Paragraph):
                    replace_in_paragraph(item, repl_func)


def make_replacer(ws):
    def _repl(text: str) -> str:
        def sub(m):
            addr, fmt = m.group(1), m.group(2)
            try:
                v = ws[addr].value
            except Exception:
                v = None
            return apply_inline_format(v, fmt)

        replaced = TOKEN_RE.sub(sub, text)

        # 간이 날짜 더미 치환
        sp = "    "
        today = datetime.today()
        today_str = f"{today.year}년{sp}{today.month}월{sp}{today.day}일"
        for token in [
            "YYYY년 MM월 DD일",
            "YYYY년    MM월    DD일",
            "YYYY 년 MM 월 DD 일",
        ]:
            replaced = replaced.replace(token, today_str)
        return replaced

    return _repl


def load_uploaded_workbook(uploaded_file) -> Workbook:
    """업로드된 엑셀을 로드하면서 친절한 오류 메시지 제공."""
    if uploaded_file is None:
        raise InvalidFileException("엑셀 파일을 업로드하세요.")

    fname = (uploaded_file.name or "").strip()
    # 확장자 체크
    if not fname.lower().endswith((".xlsx", ".xlsm")):
        raise InvalidFileException(
            f"엑셀 통합 문서(xlsx/xlsm)만 지원합니다.\n"
            f"현재 업로드된 파일: {fname}"
        )

    try:
        # ★ 여러 방법을 시도
        data = None
        
        # 방법 1: getvalue()
        try:
            data = uploaded_file.getvalue()
        except:
            pass
        
        # 방법 2: read()
        if not data or len(data) == 0:
            try:
                uploaded_file.seek(0)
                data = uploaded_file.read()
            except:
                pass
        
        # 방법 3: 직접 스트림 사용
        if not data or len(data) == 0:
            try:
                uploaded_file.seek(0)
                data = uploaded_file.getbuffer().tobytes()
            except:
                pass
        
        # 디버깅: 실제 데이터 크기 확인
        if not data or len(data) == 0:
            raise InvalidFileException(
                f"업로드된 파일이 비어있습니다.\n"
                f"파일명: {fname}\n"
                f"크기: {len(data) if data else 0} bytes\n\n"
                f"해결 방법:\n"
                f"1. 브라우저를 완전히 새로고침 (Ctrl+Shift+R 또는 Cmd+Shift+R)\n"
                f"2. 시크릿/프라이빗 모드로 접속해보세요\n"
                f"3. 파일을 다른 이름으로 저장한 후 업로드해보세요\n"
                f"4. 다른 브라우저에서 시도해보세요"
            )
        
    except InvalidFileException:
        raise
    except Exception as exc:
        raise InvalidFileException(
            f"엑셀 파일을 읽는 중 문제가 발생했습니다.\n"
            f"오류: {exc}\n"
            f"파일명: {fname}"
        ) from exc

    try:
        return load_workbook(filename=io.BytesIO(data), data_only=True)
    except BadZipFile as exc:
        raise InvalidFileException(
            "엑셀 파일이 손상되었거나 실제로는 XLS 형식일 수 있습니다.\n"
            "엑셀에서 열어서 '다른 이름으로 저장 > Excel 통합 문서 (*.xlsx)'로 다시 저장한 뒤 업로드해 보세요."
        ) from exc
    except InvalidFileException as exc:
        raise
    except Exception as exc:
        raise InvalidFileException(
            f"엑셀 파일을 여는 중 오류가 발생했습니다: {exc}"
        ) from exc


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> Optional[bytes]:
    """DOCX 바이트를 PDF 바이트로 변환(MS Word 또는 LibreOffice 필요)."""
    try:
        with tempfile.TemporaryDirectory() as td:
            in_path = os.path.join(td, "doc.docx")
            out_path = os.path.join(td, "doc.pdf")

            with open(in_path, "wb") as f:
                f.write(docx_bytes)

            # 1) docx2pdf (Windows/Office 환경)
            if docx2pdf_convert is not None:
                try:
                    docx2pdf_convert(in_path, out_path)
                    if os.path.exists(out_path):
                        with open(out_path, "rb") as f:
                            return f.read()
                except Exception:
                    pass

            # 2) LibreOffice(soffice) 사용
            if has_soffice():
                try:
                    subprocess.run(
                        [
                            "soffice",
                            "--headless",
                            "--convert-to",
                            "pdf",
                            in_path,
                            "--outdir",
                            td,
                        ],
                        check=True,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                    )
                    if os.path.exists(out_path):
                        with open(out_path, "rb") as f:
                            return f.read()
                except Exception:
                    pass
    except Exception:
        pass
    return None


def collect_leftover_tokens(doc: Document) -> Set[str]:
    """치환 후에도 남아 있는 {{...}} 토큰 수집."""
    leftovers: Set[str] = set()

    def _scan(parent):
        for item in iter_block_items(parent):
            if isinstance(item, Paragraph) and item.text:
                for m in LEFTOVER_RE.findall(item.text):
                    leftovers.add(m)

    _scan(doc)
    for section in doc.sections:
        for container in (section.header, section.footer):
            _scan(container)

    return leftovers


# ================== UI ==================
inject_style()

st.title("🧾 납입요청서 자동 생성 (DOCX + PDF)")

col_left, col_right = st.columns([1.25, 1])

with col_left:
    h4("엑셀 파일")
    # ★ key 변경으로 업로더 완전히 리셋
    if 'upload_key' not in st.session_state:
        st.session_state.upload_key = 0
    
    xlsx_file = st.file_uploader(
        "엑셀 업로드",
        type=["xlsx", "xlsm"],
        key=f"xlsx_upl_{st.session_state.upload_key}",
        help="엑셀 파일을 업로드하세요",
    )

    h4("워드 템플릿(.docx)")
    docx_tpl = st.file_uploader(
        "워드 템플릿 업로드",
        type=["docx"],
        key=f"docx_upl_{st.session_state.upload_key}",
        help="Word 템플릿 파일을 업로드하세요",
    )

    st.markdown("</div>", unsafe_allow_html=True)

    # 시트 선택은 업로드 직후 표시
    sheet_choice = None
    if xlsx_file is not None:
        # ★ 파일 크기 디버깅 정보 표시
        try:
            file_size = len(xlsx_file.getvalue())
            if file_size == 0:
                st.error(f"⚠️ 파일이 비어있습니다 (0 bytes)")
                st.info("해결 방법을 시도해보세요:")
                st.markdown("""
                1. **브라우저 새로고침** (Ctrl+Shift+R 또는 Cmd+Shift+R)
                2. **시크릿/프라이빗 창**에서 다시 접속
                3. **파일 이름을 영문으로 변경** (예: data.xlsx)
                4. **엑셀에서 다시 저장** 후 업로드
                5. **다른 브라우저** 시도 (Chrome, Edge, Firefox)
                """)
                xlsx_file = None
            else:
                small_note(f"파일 크기: {file_size:,} bytes")
        except Exception as e:
            st.warning(f"파일 정보 확인 중 오류: {e}")
        
        if xlsx_file is not None:
            try:
                wb_tmp = load_uploaded_workbook(xlsx_file)
                default_idx = (
                    wb_tmp.sheetnames.index(TARGET_SHEET)
                    if TARGET_SHEET in wb_tmp.sheetnames
                    else 0
                )
                sheet_choice = st.selectbox(
                    "사용할 시트",
                    wb_tmp.sheetnames,
                    index=default_idx,
                    key="sheet_choice",
                )
            except InvalidFileException as e:
                st.error("엑셀 파일을 읽을 수 없습니다")
                st.error(str(e))
                xlsx_file = None
            except Exception as e:
                st.warning("엑셀 미리보기 중 문제가 발생했습니다.")
                small_note(str(e))

    out_name = st.text_input("출력 파일명", value=DEFAULT_OUT)

    gen = st.button("문서 생성", use_container_width=True)

with col_right:
    st.markdown("#### 안내")
    st.markdown(
        "- **{{A1}} / {{B7|YYYY.MM.DD}} / {{C3|#,###.00}}** 형식의 인라인 포맷 지원\n"
        "- 생성 시 WORD와 PDF 제공, **개별 다운로드** 및 **ZIP 묶음** 제공\n"
        "- PDF 변환은 **MS Word(docx2pdf)** 또는 **LibreOffice(soffice)** 필요"
    )

# ================== 생성 실행 ==================
if gen:
    if not xlsx_file or not docx_tpl:
        st.error("엑셀과 템플릿을 모두 업로드하세요.")
        st.stop()

    with st.status("문서 생성 중...", expanded=True) as status:
        try:
            st.write("1) 엑셀 로드")
            wb = load_uploaded_workbook(xlsx_file)
            ws = (
                wb[sheet_choice]
                if sheet_choice
                else (
                    wb[TARGET_SHEET]
                    if TARGET_SHEET in wb.sheetnames
                    else wb[wb.sheetnames[0]]
                )
            )

            st.write("2) 템플릿 로드")
            tpl_bytes = docx_tpl.getvalue()
            doc = Document(io.BytesIO(tpl_bytes))

            st.write("3) 치환 실행")
            replacer = make_replacer(ws)
            replace_everywhere(doc, replacer)

            st.write("4) WORD 저장")
            docx_buf = io.BytesIO()
            doc.save(docx_buf)
            docx_buf.seek(0)
            docx_bytes = docx_buf.getvalue()

            st.write("5) PDF 변환 시도")
            pdf_bytes = convert_docx_to_pdf_bytes(docx_bytes)
            pdf_ok = pdf_bytes is not None

            st.write("6) 남은 토큰 확인")
            doc_after = Document(io.BytesIO(docx_bytes))
            leftovers = sorted(list(collect_leftover_tokens(doc_after)))
            if leftovers:
                with st.expander("남아 있는 토큰 목록"):
                    st.code("\n".join(leftovers))
            else:
                small_note("모든 토큰이 정상적으로 치환되었습니다.")

            status.update(label="완료", state="complete", expanded=False)
        except InvalidFileException as e:
            status.update(label="엑셀 형식 오류", state="error", expanded=True)
            st.error(str(e))
            st.stop()
        except Exception as e:
            status.update(label="오류", state="error", expanded=True)
            st.exception(e)
            st.stop()

    st.success("문서가 준비되었습니다.")

    dl_cols = st.columns(3)
    with dl_cols[0]:
        st.download_button(
            "📄 WORD 다운로드",
            data=docx_bytes,
            file_name=ensure_docx(out_name) if out_name.strip() else DEFAULT_OUT,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with dl_cols[1]:
        st.download_button(
            "🖨 PDF 다운로드",
            data=(pdf_bytes or b""),
            file_name=ensure_pdf(out_name),
            mime="application/pdf",
            disabled=not pdf_ok,
            help=None
            if pdf_ok
            else "PDF 변환 엔진(Word 또는 LibreOffice)이 없는 환경입니다.",
            use_container_width=True,
        )
    with dl_cols[2]:
        zip_buf = io.BytesIO()
        with ZipFile(zip_buf, "w", ZIP_DEFLATED) as zf:
            # WORD
            zf.writestr(
                ensure_docx(out_name) if out_name.strip() else DEFAULT_OUT,
                docx_bytes,
            )
            # PDF (가능한 경우에만)
            if pdf_ok:
                zf.writestr(ensure_pdf(out_name), pdf_bytes)
        zip_buf.seek(0)
        st.download_button(
            "📦 ZIP (WORD+PDF)",
            data=zip_buf,
            file_name=ensure_pdf(out_name).replace(".pdf", "") + "_both.zip",
            use_container_width=True,
        )
