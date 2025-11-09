"""Core document generation and formatting utilities for the Streamlit app."""

from __future__ import annotations

import io
import os
import re
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from typing import Callable, Iterable

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook

try:  # pragma: no cover - optional dependency
    from docx2pdf import convert as docx2pdf_convert
except Exception:  # pragma: no cover - best effort import
    docx2pdf_convert = None

__all__ = [
    "DEFAULT_OUT",
    "TARGET_SHEET",
    "DocumentResult",
    "collect_leftover_tokens",
    "convert_docx_to_pdf_bytes",
    "ensure_docx",
    "ensure_pdf",
    "extract_template_tokens",
    "generate_documents",
    "get_sheet_names",
]

TOKEN_RE = re.compile(r"\{\{([A-Z]+[0-9]+)(?:\|([^}]+))?\}\}")
LEFTOVER_RE = re.compile(r"\{\{[^}]+\}\}")
DEFAULT_OUT = f"{datetime.today():%Y%m%d}_#_납입요청서_DB저축은행.docx"
TARGET_SHEET = "2.  배정후 청약시"


@dataclass(slots=True)
class DocumentResult:
    """Result payload returned after generating documents."""

    docx_bytes: bytes
    pdf_bytes: bytes | None
    leftovers: list[str]


def ensure_docx(name: str) -> str:
    """Return a file name that is guaranteed to end with ``.docx``."""

    name = (name or "").strip()
    return name if name.lower().endswith(".docx") else (name + ".docx")


def ensure_pdf(name: str) -> str:
    """Return a file name that is guaranteed to end with ``.pdf``."""

    base = (name or "output").strip()
    if base.lower().endswith(".docx"):
        base = base[:-5]
    return base + ".pdf"


def has_soffice() -> bool:
    """Return ``True`` when LibreOffice's ``soffice`` binary is available."""

    return any(
        os.path.isfile(os.path.join(path, "soffice"))
        or os.path.isfile(os.path.join(path, "soffice.bin"))
        for path in os.environ.get("PATH", "").split(os.pathsep)
    )


def try_format_as_date(value) -> str:
    """Best-effort conversion of ``value`` to a formatted date string."""

    try:
        if value is None:
            return ""
        if isinstance(value, (datetime, date)):
            return f"{value.year}. {value.month}. {value.day}."
        string_value = str(value).strip()
        if re.fullmatch(r"\d{4}-\d{2}-\d{2}", string_value):
            dt_value = datetime.strptime(string_value, "%Y-%m-%d").date()
            return f"{dt_value.year}. {dt_value.month}. {dt_value.day}."
    except Exception:
        pass
    return ""


def fmt_number(value) -> str:
    """Best-effort conversion of ``value`` to a comma separated number."""

    try:
        if isinstance(value, (int, float, Decimal)):
            return f"{float(value):,.0f}"
        if isinstance(value, str):
            raw = value.replace(",", "")
            if re.fullmatch(r"-?\d+(\.\d+)?", raw):
                return f"{float(raw):,.0f}"
    except Exception:
        pass
    return ""


def value_to_text(value) -> str:
    """Convert arbitrary Excel values to a printable string."""

    as_date = try_format_as_date(value)
    if as_date:
        return as_date
    as_number = fmt_number(value)
    if as_number:
        return as_number
    return "" if value is None else str(value)


def apply_inline_format(value, fmt: str | None) -> str:
    """Apply optional inline formatting specified by ``fmt`` to ``value``."""

    if fmt is None or fmt.strip() == "":
        return value_to_text(value)

    if any(token in fmt for token in ("YYYY", "MM", "DD")):
        if isinstance(value, str) and re.fullmatch(r"\d{4}-\d{2}-\d{2}", value.strip()):
            value = datetime.strptime(value.strip(), "%Y-%m-%d").date()
        if isinstance(value, (datetime, date)):
            fmt_tokens = fmt.replace("YYYY", "%Y").replace("MM", "%m").replace("DD", "%d")
            return value.strftime(fmt_tokens)
        return value_to_text(value)

    if re.fullmatch(r"[#,0]+(?:\.[0#]+)?", fmt.replace(",", "")):
        try:
            number = float(str(value).replace(",", ""))
            decimals = 0
            if "." in fmt:
                decimals = len(fmt.split(".")[1])
            return f"{number:,.{decimals}f}"
        except Exception:
            return value_to_text(value)

    return value_to_text(value)


def iter_block_items(parent: Document | _Cell) -> Iterable[Paragraph]:
    """Yield all paragraphs from the given ``parent`` recursively."""

    if hasattr(parent, "paragraphs") and hasattr(parent, "tables"):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)


def replace_in_paragraph(paragraph: Paragraph, repl_func: Callable[[str], str]) -> None:
    """Replace the text in ``paragraph`` using ``repl_func`` preserving runs."""

    changed = False
    for run in paragraph.runs:
        new_text = repl_func(run.text)
        if new_text != run.text:
            run.text = new_text
            changed = True
    if changed:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = repl_func(full_text)
    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_everywhere(doc: Document, repl_func: Callable[[str], str]) -> None:
    """Apply ``repl_func`` to every textual element in ``doc``."""

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            replace_in_paragraph(item, repl_func)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for item in iter_block_items(container):
                if isinstance(item, Paragraph):
                    replace_in_paragraph(item, repl_func)


def make_replacer(ws) -> Callable[[str], str]:
    """Build a replacement callable that pulls values from ``ws``."""

    def _repl(text: str) -> str:
        def substitute(match: re.Match[str]) -> str:
            addr, fmt = match.group(1), match.group(2)
            try:
                value = ws[addr].value
            except Exception:
                value = None
            return apply_inline_format(value, fmt)

        replaced = TOKEN_RE.sub(substitute, text)

        separator = "    "
        today = datetime.today()
        today_str = f"{today.year}년{separator}{today.month}월{separator}{today.day}일"
        for token in ["YYYY년 MM월 DD일", "YYYY년    MM월    DD일", "YYYY 년 MM 월 DD 일"]:
            replaced = replaced.replace(token, today_str)
        return replaced

    return _repl


def convert_docx_to_pdf_bytes(docx_bytes: bytes) -> bytes | None:
    """Best-effort conversion of DOCX bytes into PDF bytes."""

    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            in_path = os.path.join(tmp_dir, "doc.docx")
            out_path = os.path.join(tmp_dir, "doc.pdf")
            with open(in_path, "wb") as doc_file:
                doc_file.write(docx_bytes)

            if docx2pdf_convert is not None:
                try:
                    docx2pdf_convert(in_path, out_path)
                    if os.path.exists(out_path):
                        with open(out_path, "rb") as pdf_file:
                            return pdf_file.read()
                except Exception:
                    pass

            if has_soffice():
                try:
                    subprocess.run(
                        ["soffice", "--headless", "--convert-to", "pdf", in_path, "--outdir", tmp_dir],
                        check=True,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                    )
                    if os.path.exists(out_path):
                        with open(out_path, "rb") as pdf_file:
                            return pdf_file.read()
                except Exception:
                    pass
    except Exception:
        pass
    return None


def collect_leftover_tokens(doc: Document) -> set[str]:
    """Collect remaining ``{{TOKEN}}`` placeholders inside ``doc``."""

    leftovers: set[str] = set()
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = "".join(run.text for run in item.runs) if item.runs else item.text
            leftovers.update(LEFTOVER_RE.findall(text or ""))
    for section in doc.sections:
        for container in (section.header, section.footer):
            for item in iter_block_items(container):
                if isinstance(item, Paragraph):
                    text = "".join(run.text for run in item.runs) if item.runs else item.text
                    leftovers.update(LEFTOVER_RE.findall(text or ""))
    return leftovers


def make_work_document(template_bytes: bytes, ws) -> Document:
    """Create a ``Document`` from ``template_bytes`` replacing tokens with ``ws`` values."""

    document = Document(io.BytesIO(template_bytes))
    replacer = make_replacer(ws)
    replace_everywhere(document, replacer)
    return document


def generate_documents(
    xlsx_bytes: bytes,
    template_bytes: bytes,
    sheet_name: str | None,
    target_sheet: str = TARGET_SHEET,
) -> DocumentResult:
    """Generate DOCX and PDF bytes along with leftover tokens."""

    workbook = load_workbook(filename=io.BytesIO(xlsx_bytes), data_only=True)
    if sheet_name and sheet_name in workbook.sheetnames:
        worksheet = workbook[sheet_name]
    else:
        worksheet = (
            workbook[target_sheet]
            if target_sheet in workbook.sheetnames
            else workbook[workbook.sheetnames[0]]
        )

    document = make_work_document(template_bytes, worksheet)
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    pdf_bytes = convert_docx_to_pdf_bytes(docx_bytes)

    document_after = Document(io.BytesIO(docx_bytes))
    leftovers = sorted(list(collect_leftover_tokens(document_after)))

    return DocumentResult(docx_bytes=docx_bytes, pdf_bytes=pdf_bytes, leftovers=leftovers)


def get_sheet_names(xlsx_bytes: bytes) -> list[str]:
    """Return the available sheet names for ``xlsx_bytes``."""

    workbook = load_workbook(filename=io.BytesIO(xlsx_bytes), data_only=True)
    return workbook.sheetnames


def extract_template_tokens(docx_bytes: bytes, *, limit: int = 12, max_paragraphs: int = 80) -> list[str]:
    """Extract up to ``limit`` placeholder tokens from the template document."""

    document = Document(io.BytesIO(docx_bytes))
    tokens: set[str] = set()
    for paragraph in document.paragraphs[:max_paragraphs]:
        tokens.update(re.findall(r"\{\{[^}]+\}\}", paragraph.text or ""))
        if len(tokens) >= limit:
            break
    return list(tokens)[:limit]
